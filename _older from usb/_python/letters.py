from docx import Document
from docx.shared import Pt

def inputdata():
    #tell tell
    print('To finish, type \'stop\'')

    #empty lists
    Names = []
    Gifts = []
    #repeat
    while True:
        #input name and gift
        Names.append(input('Name: '))
        Gifts.append(input('Gift: '))
        #check whether to stop
        if input() == 'stop':
            break

    #write results to files
    NamesFile = open('names.txt', 'w+')
    NamesFile.write(str(Names))
    GiftsFile = open('gifts.txt', 'w+')
    GiftsFile.write(str(Gifts))

    #tell of success
    print('Success')

def replace():
    #get names and gifts from files
    NamesFile = open('names.txt', 'r')
    NamesTxt = NamesFile.read()
    Names = eval(NamesTxt)
    GiftsFile = open('gifts.txt', 'r')
    GiftsTxt = GiftsFile.read()
    Gifts = eval(GiftsTxt)
    
    #get template filename
    FileName = input('Enter the name of the template file' + '\n')
    File = Document(FileName)

    #repeat for each subject
    for i in range(len(Names)):

        #create a new file named after the subject
        File.save(Names[i] + '.docx')
        NewFile = Document(Names[i] + '.docx')

        #look in each paragraph
        for paragraph in NewFile.paragraphs:
            Text = paragraph.text

            #search for replaceable terms and replace them
            if '[NAME]' in Text:
                Text = Text.replace('[NAME]', Names[i])
                paragraph.text = Text
                paragraph.style = NewFile.styles['Normal']
                
            if '[GIFT]' in Text:
                Text = Text.replace('[GIFT]', Gifts[i])
                paragraph.text = Text
                paragraph.style = NewFile.styles['Normal']
                
            if '[ADJECTIVE]' in Text:

                #find correct adjective from gift size
                if '£' in Gifts[i]:
                    Gift = Gifts[i]
                    Amount = Gift[Gift.index('£'):]
                    if ' ' in Gift:
                        Amount = Amount[1:Amount.index(' ')]
                    else:
                        Amount = Amount[1:]
                    Amount = int(Amount)
                    
                    if Amount < 10:
                        Adjective = 'nice'
                    elif Amount < 15:
                        Adjective = 'kind'
                    else:
                        Adjective = 'generous'
                else:
                    Adjective = input('What is [ADJECTIVE] of ' + Gifts[i] + '?\n')

                Text = Text.replace('[ADJECTIVE]', Adjective)
                paragraph.text = Text
                paragraph.style = NewFile.styles['Normal']

            if '[USE]' in Text:
                
                #find whether gift is money
                if '£' in Gifts[i]:
                    Use = 'save it until I find something great to buy with it'
                #if it isn't, ask user what Use is
                else:
                    Use = input('What is [USE] of ' + Gifts[i] + '?\n')

                Text = Text.replace('[USE]', Use)
                paragraph.text = Text
                paragraph.style = NewFile.styles['Normal']

        #save new file
        NewFile.save(Names[i] + '.docx')

    #tell of success
    print('Success')

#run
while True:
    #choose
    Choice = input('Type 0 to input data, 1 to create letters or \'stop\' to stop' + '\n')

    #newline
    print('')
    
    #do
    if Choice == '0':
        inputdata()
    elif Choice == '1':
        replace()
    elif Choice == 'stop':
        quit()

    #newline
    print('')
