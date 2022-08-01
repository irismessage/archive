from docx import Document

def inputdata():
    #tell tell
    print('To finish, type \'stop\'')

    #empty lists
    Names = []
    Gifts = []
    #repeat
    while True:
        #input name and gift
        Names.append(input('Name: '))
        Gifts.append(input('Gift: '))
        #check whether to stop
        if input() == 'stop':
            break

    #write results to files
    NamesFile = open('names.txt', 'w+')
    NamesFile.write(str(Names))
    GiftsFile = open('gifts.txt', 'w+')
    GiftsFile.write(str(Gifts))

    #tell of success
    print('Success')

def replace():
    #get names and gifts from files
    NamesFile = open('names.txt', 'r')
    NamesTxt = NamesFile.read()
    Names = eval(NamesTxt)
    print(Names)
    GiftsFile = open('gifts.txt', 'r')
    GiftsTxt = GiftsFile.read()
    Gifts = eval(GiftsTxt)
    print(Gifts)
    
    #get template filename
    FileName = input('Enter the name of the template file' + '\n')
    File = Document(FileName)

    #repeat for each subject
    for i in range(len(Names)):

        #create a new file named after the subject
        File.save(Names[i] + '.docx')
        NewFile = Document(Names[i] + '.docx')

        #look in each paragraph
        for paragraph in NewFile.paragraphs:
            Text = paragraph.text

            #search for replaceable terms and replace them
            if '[NAME]' in Text:
                Text = Text.replace('[NAME]', Names[i])
                paragraph.text = Text
                
            if '[GIFT]' in Text:
                Text = Text.replace('[GIFT]', Gifts[i])
                paragraph.text = Text
                
            if '[ADJECTIVE]' in Text:

                #find correct adjective from gift size
                print(Gifts[i])
                if '£' in Gifts[i]:
                    Gift = Gifts[i]
                    print(Gift)
                    Amount = Gift[Gift.index('£'):]
                    print(Amount)
                    if ' ' in Gift:
                        print('space')
                        Amount = Amount[1:Amount.index(' ')]
                    else:
                        print('no space')
                        Amount = Amount[1:]
                    print(Amount)
                    Amount = int(Amount)
                    
                    if Amount < 10:
                        Adjective = 'nice'
                    elif Amount < 15:
                        Adjective = 'kind'
                    else:
                        Adjective = 'generous'
                else:
                    Adjective = 'kind'

                Text = Text.replace('[ADJECTIVE]', Adjective)
                paragraph.text = Text

        #save new file     
        NewFile.save(Names[i] + '.docx')

    #tell of success
    print('Success')

#run
while True:
    #choose
    Choice = input('Type 0 to input data, 1 to create letters or \'stop\' to stop' + '\n')

    #newline
    print('')
    
    #do
    if Choice == '0':
        inputdata()
    elif Choice == '1':
        replace()
    elif Choice == 'stop':
        quit()

    #newline
    print('')
